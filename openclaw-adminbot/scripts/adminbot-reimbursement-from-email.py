#!/usr/bin/env python3
import base64
import io
import json
import os
import re
import sys
from copy import deepcopy
from datetime import datetime
from pathlib import Path

import pypdfium2
from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

COMPUTE_EXPENSE_TEMPLATE = "Compute_Expense_Form.xlsx"
TRIP_SUMMARY_TEMPLATE = "Trip_Summary_Form.docx"

USER_FORMS = Path.home() / ".openclaw" / "skills" / "adminbot-reimbursements" / "forms"
# The same templates ship inside the repo, so a deployed release always carries a usable copy.
BUNDLED_FORMS = (
    Path(__file__).resolve().parent.parent
    / "extensions"
    / "adminbot"
    / "skills"
    / "adminbot-reimbursements"
    / "forms"
)


def resolve_forms_dir() -> Path:
    """Prefer host-local templates, then the ones bundled with this release.

    The user copy wins when present so a host can carry an updated or per-claimant form without a
    redeploy. Falling back to the bundled copy is what keeps a fresh deployment working: the forms
    live outside the repo checkout on the dev box, so a new host would otherwise fail at generation
    time with a missing-template error even though the release shipped the very same files.
    """
    override = os.environ.get("ADMINBOT_REIMBURSEMENT_FORMS_DIR")
    if override:
        return Path(override)
    if (USER_FORMS / COMPUTE_EXPENSE_TEMPLATE).is_file():
        return USER_FORMS
    return BUNDLED_FORMS


FORMS = resolve_forms_dir()

PDF_RENDER_MAX_PAGES = 6
PDF_RENDER_DPI = 150

def extract_receipt(path: Path) -> dict:
    """Pull both a text layer and page images out of a receipt PDF.

    pypdfium2 ships PDFium as a self-contained wheel, so this works on hosts without root (Aurora
    included); the previous poppler `pdftotext`/`pdftoppm` shell-outs silently produced nothing at
    all when those binaries were absent, which read downstream as "the user's receipts are
    unreadable". Page images matter independently of text: scanned receipts carry no text layer, so
    the vision model is the only way to read them.
    """
    text_parts: list[str] = []
    images: list[dict] = []
    try:
        pdf = pypdfium2.PdfDocument(path)
    except Exception as error:  # pypdfium2 raises PdfiumError for corrupt/encrypted input
        print(f"warning: {path.name}: cannot open PDF: {error}", file=sys.stderr)
        return {
            "name": path.name,
            "text": f"[{path.name}: file could not be opened as a PDF; retain as supporting attachment]",
            "images": [],
        }
    try:
        for index in range(min(len(pdf), PDF_RENDER_MAX_PAGES)):
            page = pdf[index]
            text_page = page.get_textpage()
            text_parts.append(text_page.get_text_range())
            bitmap = page.render(scale=PDF_RENDER_DPI / 72)
            buffer = io.BytesIO()
            bitmap.to_pil().save(buffer, format="PNG")
            images.append(
                {
                    "media_type": "image/png",
                    "data_base64": base64.b64encode(buffer.getvalue()).decode("ascii"),
                }
            )
    finally:
        pdf.close()
    text = "\n".join(text_parts)
    if not text.strip() and not images:
        text = f"[{path.name}: unable to extract text or render pages; retain as supporting attachment]"
    return {"name": path.name, "text": text, "images": images}


def require_template(name: str) -> Path:
    template = FORMS / name
    if not template.is_file():
        # Naming every candidate turns "which of these paths was I supposed to populate?" into a
        # one-line answer, which is exactly what the bare path failed to convey on a fresh host.
        raise FileNotFoundError(
            f"required reimbursement template is missing: {template} "
            f"(looked in {FORMS}; set ADMINBOT_REIMBURSEMENT_FORMS_DIR, or place it in "
            f"{USER_FORMS}, or ship it at {BUNDLED_FORMS})"
        )
    return template


def excel_date(value):
    if not value:
        return None
    for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y", "%Y/%m/%d"):
        try:
            return datetime.strptime(str(value), fmt).date()
        except ValueError:
            pass
    return str(value)


CATEGORY_COLUMNS = {
    "hotel": {"canada": "K", "usa": "L", "international": "L"},
    "accommodation": {"canada": "K", "usa": "L", "international": "L"},
    "rail": {"canada": "M", "usa": "N", "international": "N"},
    "bus": {"canada": "M", "usa": "N", "international": "N"},
    "transit": {"canada": "O", "usa": "O", "international": "O"},
    "car rental": {"canada": "P", "usa": "P", "international": "P"},
    "meal": {"canada": "Q", "usa": "S", "international": "S"},
    "food": {"canada": "Q", "usa": "S", "international": "S"},
    "taxi": {"canada": "T", "usa": "V", "international": "V"},
    "rideshare": {"canada": "T", "usa": "V", "international": "V"},
    "uber": {"canada": "T", "usa": "V", "international": "V"},
    "lyft": {"canada": "T", "usa": "V", "international": "V"},
    "conference": {"canada": "W", "usa": "W", "international": "W"},
    "registration": {"canada": "W", "usa": "W", "international": "W"},
    "hospitality": {"canada": "X", "usa": "X", "international": "X"},
    "parking": {"canada": "Z", "usa": "Z", "international": "Z"},
}


def normalized_category(category: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", category.strip().lower()).strip()


def category_column(expense: dict) -> str:
    normalized = normalized_category(str(expense.get("category") or "other"))
    region = str(expense.get("region") or "international").lower()
    if "airfare" in normalized or "flight" in normalized or "air travel" in normalized:
        columns = (
            {"canada": "H", "usa": "I", "international": "J"}
            if expense.get("airfare_class") == "above_economy"
            else {"canada": "E", "usa": "F", "international": "G"}
        )
        return columns.get(region, columns["international"])
    for key, columns in CATEGORY_COLUMNS.items():
        if key in normalized:
            return columns.get(region, columns["international"])
    return "AA"


def requested_currency(data: dict) -> str:
    currency = str(data.get("currency") or "CAD").upper()
    if currency == "OTHER":
        return str(data.get("other_currency") or "OTHER").upper()
    return currency


# Every cell the claimant is expected to touch on the "Reimbursement Form" tab is highlighted
# yellow in the template. Everything else there -- the L "AMOUNT" column, its total rows, and the
# G/L, tax-code, and cost-centre fields -- sits under "TO BE COMPLETED BY BUSINESS OFFICER", so the
# generated form leaves it untouched rather than pre-empting the business officer's entries.
YELLOW_CURRENCY_CHECKBOXES = {"CAD": "I4", "USD": "I5"}
OTHER_CURRENCY_CHECKBOX = "I6"
OTHER_CURRENCY_LABEL_CELL = "I7"

# Only these three categories have a dedicated TIP column on the summary sheet (meals, taxi, and
# hospitality, Canadian columns only). Anything else has nowhere to put a tip, so it stays inside
# the expense amount.
SUMMARY_TIP_COLUMNS = {"Q": "R", "T": "U", "X": "Y"}
SUMMARY_FIRST_ROW = 4
SUMMARY_LAST_ROW = 33
SUMMARY_TOTAL_ROW = 34
SUMMARY_GRAND_TOTAL_CELL = "D36"


def summary_description(expense: dict) -> str:
    """Receipts are frequently in a currency other than the claim currency, and the sheet has no
    currency column, so the exact currency rides along in DESCRIPTION. Amounts are never converted.
    """
    description = str(expense.get("description") or expense.get("category") or "Expense").strip()
    code = str(expense.get("currency") or "").strip().upper()
    if not code:
        return description
    return f"{description} ({code} {float(expense.get('amount') or 0):,.2f})"


def fill_workbook(data: dict, output: Path) -> None:
    wb = load_workbook(require_template(COMPUTE_EXPENSE_TEMPLATE))
    main = wb["Reimbursement Form"]
    main["A10"] = data.get("personnel_number") or ""
    main["B10"] = data.get("travel_period") or data.get("trip_dates") or ""
    main["A12"] = data.get("claimant_name") or ""
    main["A14"] = data.get("claimant_address") or ""
    main["A18"] = data.get("purpose") or ""
    main["A38"] = data.get("claimant_name") or ""
    main["C38"] = data.get("claimant_title") or ""
    currency = requested_currency(data)
    for cell in (*YELLOW_CURRENCY_CHECKBOXES.values(), OTHER_CURRENCY_CHECKBOX, OTHER_CURRENCY_LABEL_CELL):
        main[cell] = None
    if currency in YELLOW_CURRENCY_CHECKBOXES:
        main[YELLOW_CURRENCY_CHECKBOXES[currency]] = "X"
    else:
        main[OTHER_CURRENCY_CHECKBOX] = "X"
        main[OTHER_CURRENCY_LABEL_CELL] = currency

    summary = wb["Expense Summary"]
    other_labels: dict[str, str] = {}
    for index, expense in enumerate(
        data.get("expenses", [])[: SUMMARY_LAST_ROW - SUMMARY_FIRST_ROW + 1], start=SUMMARY_FIRST_ROW
    ):
        summary[f"B{index}"] = expense.get("receipt_number") or index - SUMMARY_FIRST_ROW + 1
        summary[f"C{index}"] = excel_date(expense.get("date"))
        summary[f"D{index}"] = summary_description(expense)
        column = category_column(expense)
        if column == "AA":
            label = str(expense.get("category") or "Other")[:30]
            column = other_labels.setdefault(label, ("AA", "AB", "AC", "AD")[min(len(other_labels), 3)])
            summary[f"{column}2"] = label
        amount = float(expense.get("amount") or 0)
        tip = float(expense.get("tip_amount") or 0)
        tip_column = SUMMARY_TIP_COLUMNS.get(column)
        if tip_column and 0 < tip < amount:
            # Split so the two columns still add back up to the receipt total. A tip at or above the
            # total is a misread rather than a real tip, so it falls through to the inline form
            # instead of writing a zero or negative base.
            summary[f"{column}{index}"] = round(amount - tip, 2)
            summary[f"{tip_column}{index}"] = tip
        else:
            # No itemized tip, or no TIP column for this category: the form accepts a tip-inclusive
            # amount directly (its "ON (13%HST+TIP)" rows exist for exactly this).
            summary[f"{column}{index}"] = amount
    # Restated rather than assumed: a claim is only useful if the totals are actually there.
    for column_index in range(5, 31):
        column = get_column_letter(column_index)
        summary[f"{column}{SUMMARY_TOTAL_ROW}"] = (
            f"=SUM({column}{SUMMARY_FIRST_ROW}:{column}{SUMMARY_LAST_ROW})"
        )
    summary[SUMMARY_GRAND_TOTAL_CELL] = f"=SUM(E{SUMMARY_TOTAL_ROW}:AD{SUMMARY_TOTAL_ROW})"
    calculation = getattr(wb, "calculation", None)
    if calculation is not None:
        calculation.fullCalcOnLoad = True
        calculation.forceFullCalc = True
    wb.save(output)


def replace_paragraph(paragraph, text: str) -> None:
    run_properties = (
        deepcopy(paragraph.runs[0]._r.rPr)
        if paragraph.runs and paragraph.runs[0]._r.rPr is not None
        else None
    )
    paragraph_properties = paragraph._p.pPr
    for child in list(paragraph._p):
        if paragraph_properties is None or child is not paragraph_properties:
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)


def set_cell_text(cell, text: str) -> None:
    replace_paragraph(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        replace_paragraph(paragraph, "")


# The template's expense table ships five fixed rows in this order, then a TOTAL row. Their labels
# (and the taxi row's pick-up/drop-off instructions) are part of the form, so the generated document
# fills only the amount cells and leaves every label exactly as the template wrote it.
TRIP_ROWS = (
    ("Hotel", ("hotel", "accommodation", "lodging")),
    ("Flight", ("airfare", "flight", "air travel")),
    ("Meals", ("meal", "food", "per diem", "allowance")),
    (
        "Taxi Expenses",
        ("taxi", "rideshare", "uber", "lyft", "rail", "bus", "transit", "car rental", "parking", "mileage"),
    ),
    ("Conference Registration", ("conference", "registration")),
)

# What the claimant must physically attach for each row, worded as the template's own Enclosed
# section words it.
TRIP_ENCLOSURES = {
    "Hotel": "HOTEL: Receipt + Proof of Stay",
    "Flight": "FLIGHT: Receipts + Boarding Pass / Confirmation",
    "Meals": "MEALS: Itemized Receipts",
    "Taxi Expenses": "TAXI: Receipts + CC Statement",
    "Conference Registration": "Conference: Receipt + CC Statement",
}

# CAD and USD both render as "$", which on a reimbursement claim is a real ambiguity rather than a
# cosmetic one, so the two dollars stay distinguishable.
CURRENCY_SYMBOLS = {
    "CAD": "C$",
    "USD": "US$",
    "EUR": "\u20ac",
    "GBP": "\u00a3",
    "JPY": "\u00a5",
    "CHF": "CHF ",
    "AUD": "A$",
    "NZD": "NZ$",
}


def money(amount: float, currency: str) -> str:
    code = (currency or "").strip().upper()
    symbol = CURRENCY_SYMBOLS.get(code)
    return f"{symbol}{amount:,.2f}" if symbol else f"{code + ' ' if code else ''}{amount:,.2f}"


def trip_row_label(expense: dict) -> str | None:
    normalized = normalized_category(str(expense.get("category") or ""))
    for label, keywords in TRIP_ROWS:
        if any(keyword in normalized for keyword in keywords):
            return label
    return None


def trip_expense_rows(data: dict) -> tuple[dict[str, list[dict]], list[str]]:
    """Group expenses under the template's fixed rows, keeping each expense separate.

    Expenses that match none of the fixed rows keep their own category as an extra row rather than
    being folded into an existing one, which would misfile them.
    """
    grouped: dict[str, list[dict]] = {label: [] for label, _ in TRIP_ROWS}
    extra_order: list[str] = []
    for expense in data.get("expenses", []):
        label = trip_row_label(expense)
        if label is None:
            label = str(expense.get("category") or "Other expenses").strip() or "Other expenses"
            if label not in grouped:
                grouped[label] = []
                extra_order.append(label)
        grouped[label].append(expense)
    return grouped, extra_order


def trip_row_text(expenses: list[dict], claim_currency: str) -> str:
    """One entry per expense, joined with " + " as the reference form does, so a row reads
    "C$20.00 (lunch at X) + C$30.00 (dinner at Y)" instead of collapsing to a single number."""
    if not expenses:
        return money(0, claim_currency)
    parts = []
    for expense in expenses:
        amount = float(expense.get("amount") or 0)
        rendered = money(amount, str(expense.get("currency") or claim_currency))
        description = str(expense.get("description") or "").strip()
        parts.append(f"{rendered} ({description})" if description else rendered)
    return " + ".join(parts)


def trip_total_text(data: dict, claim_currency: str) -> str:
    """Totals stay per-currency and are never converted, matching the reference form's
    "\u00a3558 + \u20ac180.85 + $70.17" style."""
    totals: dict[str, float] = {}
    for expense in data.get("expenses", []):
        code = str(expense.get("currency") or claim_currency).strip().upper()
        totals[code] = totals.get(code, 0.0) + float(expense.get("amount") or 0)
    if not totals:
        return money(0, claim_currency)
    return " + ".join(money(amount, code) for code, amount in totals.items())


def append_table_row(table):
    """Clone the last expense row so an extra category keeps the template's borders and spacing."""
    new_row = deepcopy(table.rows[-2]._tr)
    table.rows[-1]._tr.addprevious(new_row)
    return table.rows[-2]


def strip_example_section(doc) -> None:
    """The template carries a worked EXAMPLE after the real form. It is reference material for
    filling the document, never part of a submitted claim, so everything from that heading onward
    (including the example's own expense table) is dropped."""
    body = doc.element.body
    marker = next(
        (p._p for p in doc.paragraphs if p.text.strip().upper().startswith("EXAMPLE")), None
    )
    if marker is None:
        return
    removing = False
    for child in list(body):
        if child is marker:
            removing = True
        if removing and not child.tag.endswith("}sectPr"):
            body.remove(child)


def set_enclosure_lines(doc, lines: list[str], first_index: int, last_index: int) -> None:
    """Write the Enclosed list, reusing the template's four paragraphs and cloning that formatting
    when a claim needs more than four."""
    for offset, index in enumerate(range(first_index, last_index + 1)):
        replace_paragraph(doc.paragraphs[index], lines[offset] if offset < len(lines) else "")
    template_paragraph = doc.paragraphs[first_index]._p
    anchor = doc.paragraphs[last_index]._p
    for line in lines[last_index - first_index + 1 :]:
        clone = deepcopy(template_paragraph)
        anchor.addnext(clone)
        anchor = clone
        replace_paragraph(
            next(p for p in doc.paragraphs if p._p is clone),
            line,
        )


def remove_paragraphs(doc, indices: list[int]) -> None:
    for paragraph in [doc.paragraphs[index] for index in indices if index < len(doc.paragraphs)]:
        paragraph._p.getparent().remove(paragraph._p)


def fill_trip_summary(data: dict, output: Path) -> None:
    doc = Document(require_template(TRIP_SUMMARY_TEMPLATE))
    strip_example_section(doc)
    values = {
        1: data.get("trip_title") or "Travel reimbursement",
        2: data.get("trip_dates") or data.get("travel_period") or "",
        3: data.get("trip_location") or "",
        10: data.get("claimant_name") or "",
        11: data.get("claimant_email") or "",
        13: data.get("claimant_title") or "",
        16: data.get("purpose") or "",
    }
    for index, value in values.items():
        if index < len(doc.paragraphs):
            replace_paragraph(doc.paragraphs[index], str(value))
    claim_currency = requested_currency(data)
    grouped, extra_order = trip_expense_rows(data)
    if doc.tables:
        table = doc.tables[0]
        for index, (label, _) in enumerate(TRIP_ROWS):
            set_cell_text(table.rows[index].cells[1], trip_row_text(grouped[label], claim_currency))
        for label in extra_order:
            row = append_table_row(table)
            set_cell_text(row.cells[0], label)
            set_cell_text(row.cells[1], trip_row_text(grouped[label], claim_currency))
        set_cell_text(table.rows[-1].cells[1], trip_total_text(data, claim_currency))
    enclosures = [
        TRIP_ENCLOSURES[label] for label, _ in TRIP_ROWS if grouped[label]
    ]
    enclosures += [f"{label.upper()}: Receipt" for label in extra_order]
    set_enclosure_lines(doc, enclosures, 23, 26)
    # The template reserves a second event block (paragraphs 5-7) after the first. A draft describes
    # one trip, so those lines are removed rather than left blank, which would leave five blank
    # lines where the reference form has one.
    remove_paragraphs(doc, [8, 7, 6, 5])
    doc.save(output)


def main():
    if len(sys.argv) < 3:
        raise SystemExit("usage: helper extract FILE... | fill INPUT_JSON OUTPUT_DIR")
    if sys.argv[1] == "extract":
        receipts = [extract_receipt(Path(item)) for item in sys.argv[2:]]
        print(json.dumps({"receipts": receipts}))
        return
    if sys.argv[1] != "fill" or len(sys.argv) != 4:
        raise SystemExit("usage: helper fill INPUT_JSON OUTPUT_DIR")
    data = json.loads(Path(sys.argv[2]).read_text())
    output_dir = Path(sys.argv[3])
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", str(data.get("claimant_name") or "claimant")).strip("_")
    workbook = output_dir / f"Compute_Expense_Form_{safe_name}.xlsx"
    trip = output_dir / f"Trip_Summary_Form_{safe_name}.docx"
    fill_workbook(data, workbook)
    fill_trip_summary(data, trip)
    print(json.dumps({"files": [str(workbook), str(trip)]}))


if __name__ == "__main__":
    main()
